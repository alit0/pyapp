import base64
import io
from typing import Optional
import PyPDF2
from docx import Document
import openpyxl

class FileProcessor:
    """Clase para procesar diferentes tipos de archivos y extraer su contenido como texto."""
    
    @staticmethod
    def decode_base64_file(base64_content: str) -> bytes:
        """Decodifica el contenido base64 del archivo."""
        print("🔄 Decodificando archivo base64...")
        print(f"📏 Longitud del contenido base64: {len(base64_content)}")
        
        # Remover el prefijo data:tipo/subtipo;base64, si existe
        if ',' in base64_content:
            prefix, base64_data = base64_content.split(',', 1)
            print(f"🏷️  Prefijo detectado: {prefix}")
            print(f"📏 Longitud después de remover prefijo: {len(base64_data)}")
            base64_content = base64_data
        
        try:
            decoded = base64.b64decode(base64_content)
            print(f"✅ Archivo decodificado exitosamente ({len(decoded)} bytes)")
            
            # Verificar que los primeros bytes sean correctos para diferentes formatos
            if len(decoded) >= 4:
                first_bytes = decoded[:4]
                print(f"🔍 Primeros 4 bytes (hex): {first_bytes.hex()}")
                
                # Verificar signatura de archivo
                if first_bytes == b'PK\x03\x04':
                    print("✅ Signatura ZIP/DOCX/XLSX detectada")
                elif first_bytes == b'%PDF':
                    print("✅ Signatura PDF detectada")
                else:
                    print(f"⚠️  Signatura desconocida: {first_bytes}")
            
            return decoded
        except Exception as e:
            print(f"❌ Error en decodificación base64: {str(e)}")
            raise e
    
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extrae texto de un archivo PDF."""
        try:
            print("📖 Extrayendo texto de PDF...")
            print(f"📏 Tamaño del archivo: {len(file_bytes)} bytes")
            
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            print(f"📄 PDF tiene {len(pdf_reader.pages)} páginas")
            
            text = ""
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text += page_text + "\n"
                print(f"  - Página {i+1}: {len(page_text)} caracteres extraídos")
            
            print(f"✅ Extracción de PDF completada ({len(text)} caracteres totales)")
            return text.strip()
        except Exception as e:
            error_msg = f"Error al leer PDF: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg
    
    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extrae texto de un archivo DOCX."""
        try:
            print("📝 Extrayendo texto de DOCX...")
            print(f"📏 Tamaño del archivo: {len(file_bytes)} bytes")
            
            # Verificar que el archivo tenga la signatura correcta
            if len(file_bytes) >= 4:
                signature = file_bytes[:4]
                print(f"🔍 Signatura del archivo: {signature.hex()}")
                if signature != b'PK\x03\x04':
                    print(f"⚠️  ADVERTENCIA: El archivo no tiene signatura ZIP/DOCX válida")
            
            docx_file = io.BytesIO(file_bytes)
            
            try:
                doc = Document(docx_file)
                print(f"📄 DOCX cargado exitosamente, tiene {len(doc.paragraphs)} párrafos")
            except Exception as doc_error:
                print(f"❌ Error al cargar documento DOCX: {str(doc_error)}")
                
                # Intentar diagnóstico adicional
                docx_file.seek(0)
                first_100_bytes = docx_file.read(100)
                print(f"🔍 Primeros 100 bytes del archivo: {first_100_bytes[:50].hex()}...")
                
                # Intentar verificar si es un archivo ZIP válido
                import zipfile
                docx_file.seek(0)
                try:
                    with zipfile.ZipFile(docx_file, 'r') as zip_file:
                        print(f"✅ Archivo ZIP válido, contiene: {zip_file.namelist()[:5]}")
                except zipfile.BadZipFile as zip_error:
                    print(f"❌ No es un archivo ZIP válido: {str(zip_error)}")
                
                raise doc_error
            
            text = ""
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                if i < 5:  # Solo mostrar los primeros 5 párrafos
                    print(f"  - Párrafo {i+1}: {len(paragraph.text)} caracteres")
            
            print(f"✅ Extracción de DOCX completada ({len(text)} caracteres totales)")
            return text.strip()
        except Exception as e:
            error_msg = f"Error al leer DOCX: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg
    
    @staticmethod
    def extract_text_from_xlsx(file_bytes: bytes) -> str:
        """Extrae texto de un archivo XLSX."""
        try:
            print("📊 Extrayendo datos de XLSX...")
            print(f"📏 Tamaño del archivo: {len(file_bytes)} bytes")
            
            xlsx_file = io.BytesIO(file_bytes)
            workbook = openpyxl.load_workbook(xlsx_file)
            
            print(f"📄 XLSX tiene {len(workbook.sheetnames)} hojas: {workbook.sheetnames}")
            
            text = ""
            total_rows = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n=== HOJA: {sheet_name} ===\n"
                
                # Obtener todas las filas con datos
                rows_with_data = []
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell).strip())
                    if any(cell for cell in row_text if cell):  # Solo si hay contenido
                        rows_with_data.append(row_text)
                
                print(f"  - Hoja '{sheet_name}': {len(rows_with_data)} filas con datos")
                
                # Agregar cada fila al texto
                for i, row_data in enumerate(rows_with_data):
                    if i == 0:
                        # Primera fila como encabezados
                        text += "ENCABEZADOS: " + " | ".join(row_data) + "\n"
                    else:
                        # Resto de filas como datos
                        text += f"FILA {i}: " + " | ".join(row_data) + "\n"
                    total_rows += 1
                
                text += f"\n--- FIN HOJA {sheet_name} ({len(rows_with_data)} filas) ---\n"
            
            print(f"✅ Extracción de XLSX completada:")
            print(f"  📊 Total de filas procesadas: {total_rows}")
            print(f"  📄 Caracteres totales: {len(text)}")
            
            return text.strip()
        except Exception as e:
            error_msg = f"Error al leer XLSX: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg
    
    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        """Extrae texto de un archivo TXT."""
        try:
            print("📄 Extrayendo texto de TXT...")
            print(f"📏 Tamaño del archivo: {len(file_bytes)} bytes")
            
            # Intentar diferentes codificaciones
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_bytes.decode(encoding)
                    print(f"✅ TXT decodificado con codificación '{encoding}' ({len(text)} caracteres)")
                    return text
                except UnicodeDecodeError:
                    print(f"  - Fallo con codificación '{encoding}', probando siguiente...")
                    continue
            
            error_msg = "Error: No se pudo decodificar el archivo de texto"
            print(f"❌ {error_msg}")
            return error_msg
        except Exception as e:
            error_msg = f"Error al leer TXT: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg
    
    @classmethod
    def process_file(cls, base64_content: str, file_type: str, file_name: str) -> str:
        """
        Procesa un archivo según su tipo y retorna el texto extraído.
        
        Args:
            base64_content: Contenido del archivo en base64
            file_type: Tipo MIME del archivo
            file_name: Nombre del archivo
        
        Returns:
            Texto extraído del archivo
        """
        try:
            print("=== PROCESANDO ARCHIVO ===")
            print(f"📄 Archivo: {file_name}")
            print(f"🏷️  Tipo MIME: {file_type}")
            print(f"📏 Longitud contenido base64: {len(base64_content)}")
            
            file_bytes = cls.decode_base64_file(base64_content)
            
            # Determinar el tipo de archivo por extensión o MIME type
            file_extension = file_name.lower().split('.')[-1] if '.' in file_name else ''
            print(f"🔍 Extensión detectada: .{file_extension}")
            
            if file_extension == 'pdf' or 'pdf' in file_type.lower():
                print("📖 Procesando como PDF...")
                return cls.extract_text_from_pdf(file_bytes)
            elif file_extension == 'docx' or 'wordprocessingml' in file_type.lower():
                print("📝 Procesando como DOCX...")
                return cls.extract_text_from_docx(file_bytes)
            elif file_extension in ['xlsx', 'xls'] or 'spreadsheet' in file_type.lower():
                print("📊 Procesando como XLSX...")
                return cls.extract_text_from_xlsx(file_bytes)
            elif file_extension == 'txt' or 'text/plain' in file_type.lower():
                print("📄 Procesando como TXT...")
                return cls.extract_text_from_txt(file_bytes)
            else:
                error_msg = f"Tipo de archivo no soportado: {file_type} (.{file_extension})"
                print(f"❌ {error_msg}")
                return error_msg
                
        except Exception as e:
            error_msg = f"Error al procesar el archivo {file_name}: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg